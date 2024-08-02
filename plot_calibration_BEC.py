from pathlib import Path
import matplotlib.pylab as plt
import json
#from ngen_cal import plot_objective, plot_stuff, plot_obs, plot_output, plot_parameter_space
import pandas as pd
import os
import sys
sys.path.append("/home/west/Projects/CAMELS/NextGenCode/Model_evaluation")
import objectives as OB
import numpy as np
import shutil
from datetime import datetime


#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Fri Jul 29 14:01:06 2022

@author: west
"""
from pathlib import Path
import matplotlib.pylab as plt
import json
#from ngen_cal import plot_objective, plot_stuff, plot_obs, plot_output, plot_parameter_space
import pandas as pd
import os,glob
import subprocess
import time
import yaml
import seaborn as sns
import geopandas as gp
from hydrotools.metrics import metrics 
import math as m

# BChoat: Seems like very few of these pptx/docx imports are actually used
import pptx
from pptx.chart.data import CategoryChartData
from pptx.enum.chart import XL_CHART_TYPE
from pptx.enum.text import PP_ALIGN
from pptx.util import Inches
from pptx import Presentation
#import df2img
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION


def change_orientation2():
    current_section = document.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    # new_section = document.add_section(WD_SECTION.NEW_PAGE)
    current_section.orientation = WD_ORIENT.LANDSCAPE
    current_section.page_width = new_width
    current_section.page_height = new_height

    return current_section

def plot_calibration_1(Model_name,Q_cms,Runoff_m3s,Next_Total_out,NWM_21,Objective,total_area,output_figure):
    
    plt.rc('axes', titlesize=12) 
    
    Q_cms.index.names = ['time']
    # Q_cms_stats=pd.concat([Q_cms,NWM_21],axis=1).dropna()
    # print('Q_cms')
    #print(Q_cms)
    # print('NWM_21')
    # print(NWM_21)

    Q_cms_stats = pd.merge(Q_cms, NWM_21, 
                     left_index=True, right_index=True,
                     how = 'left').dropna()

    KGE=metrics.kling_gupta_efficiency(Q_cms_stats['q_cms_obs'].values, Q_cms_stats['q_cms_sim'].values)
    # print(f"q_cms_obs: {Q_cms_stats['q_cms_obs'].values}; q_cms_sim: {Q_cms_stats['q_cms_sim'].values}")
    # print(f'KGE: {KGE}')
    Nash=metrics.nash_sutcliffe_efficiency(Q_cms_stats['q_cms_obs'].values, Q_cms_stats['q_cms_sim'].values)
    NNash=1./(2.-Nash)  
    
    # print(f'NWM_21: {NWM_21}')
    if(len(NWM_21)>0):          
        Nash_NWM=metrics.nash_sutcliffe_efficiency(Q_cms_stats['q_cms_obs'].values, Q_cms_stats['flow_cms'].values)
        NNash_NWM=1./(2.-Nash_NWM) 
        KGE_NWM=metrics.kling_gupta_efficiency(Q_cms_stats['q_cms_obs'].values, Q_cms_stats['flow_cms'].values)
    else: 
        Nash_NWM=""
        NNash_NWM=""
        KGE_NWM=""
    
    #mean_bias=metrics.mean_error(Q_cms['q_cms_obs'], Q_cms[['q_cms_sim']])
    #mean_absolute_error=OB.mean_absolute_error(Q_cms['q_cms_obs'], Q_cms[['q_cms_sim']])
    
    #RMSE=metrics.root_mean_squared_error(Q_cms['q_cms_obs'], Q_cms[['q_cms_sim']])
    #NRMSE = sqrt(RMSE)/Total_dailyNoNAN['Obs_Q'].mean()
    #NRMSE=OB.NRMSE(Total_dailyNoNAN[Var_name_out],Total_dailyNoNAN['Obs_Q'])
    
    NWM_21_Acum=NWM_21.cumsum(axis= 0)*3.600/(total_area*1000)
    Q_cms_Acum=Q_cms.cumsum(axis= 0)*3.600/(total_area*1000)
    if(len(Runoff_m3s)>0): Runoff_m3s_Acum=Runoff_m3s.cumsum(axis= 0)
    if(len(Next_Total_out)>0): Next_Total_out_Acum=Next_Total_out.cumsum(axis= 0)
       
    num_subplots = 3 #If plotting cumulative precipitation/soil moisture, use 3 subplots (flow, cumulative rain, soil moisture)
    plt.rcParams["figure.figsize"] = (8,8)
    fig, ax_compare = plt.subplots(num_subplots, 1) #Create the plot for this zone
    #ax_compare[0].set_title('Cumulative Precipitation')
    
    if(len(NWM_21)>0): ax_compare[0].plot (NWM_21.index,NWM_21_Acum['flow_cms'], label = 'NWM 2.1 runoff', color= 'forestgreen',linewidth=2.5) #Plot obs data on comparison figure
    ax_compare[0].plot (Q_cms_Acum.index,Q_cms_Acum['q_cms_obs'],'o',markersize=2, label = 'Obs discharge', color= 'blue') #Plot obs data on comparison figure
    if('q_cms_sim' in Q_cms_Acum.columns): ax_compare[0].plot (Q_cms_Acum.index,Q_cms_Acum['q_cms_sim'], label = Model_name, color= 'orange',linewidth=2.5) #Plot obs data on comparison figure    
    Model_name_NOAH="NOM_"+Model_name.split("_")[1]
    if('q_cms_NOAH_sim' in Q_cms_Acum.columns) & ("PET" in Model_name): ax_compare[0].plot (Q_cms_Acum.index,Q_cms_Acum['q_cms_NOAH_sim'], label = Model_name_NOAH, color= 'orange',linewidth=2.5,linestyle="dashed") #Plot obs data on comparison figure    
    
    if(len(Runoff_m3s)>0): ax_compare[0].plot (Runoff_m3s_Acum.index,Runoff_m3s_Acum, label = Model_name+'-cat', color= 'mediumpurple',linewidth=2) #Plot obs data on comparison figure
    if(len(Next_Total_out)>0): ax_compare[0].plot (Next_Total_out.index,Next_Total_out_Acum, label = Model_name+'-nex', color= 'salmon',linewidth=1.2)
    ax_compare[0].set_xlim([Q_cms.index.min(),Q_cms.index.max()])
    
    #ax_compare[0].plot (Total.index,Total_Acum['Rainfall_from_forcing'],label = 'Precipitation', color= 'silver') #Plot obs data on comparison figure
    ax_compare[0].legend(loc='upper left')
    #ax_compare[0].set_xlabel ('Date'); 
    ax_compare[0].set_ylabel('Cumulative Q (m)')

    
    if(len(NWM_21)>0): ax_compare[1].plot (NWM_21.index,NWM_21['flow_cms'], label = 'NWM 2.1 runoff'+" (KGE={:.2f}".format(KGE_NWM)+")", color= 'forestgreen',linewidth=2.5) #Plot obs data on comparison figure            
    ax_compare[1].plot(Q_cms.index,Q_cms['q_cms_obs'],'o',markersize=2, label = 'Obs runoff', color= 'blue') #Plot obs data on comparison figure
    if('q_cms_sim' in Q_cms_Acum.columns): ax_compare[1].plot(Q_cms.index,Q_cms['q_cms_sim'], label = Model_name+" (KGE={:.2f}".format(KGE)+")", color= 'orange',linewidth=1.2) #Plot obs data on comparison figure   
 
    ax_compare[1].set_xlim([Q_cms.index.min(),Q_cms.index.max()])
    if(len(Next_Total_out)>0): ax_compare[1].plot(Next_Total_out.index,Next_Total_out, label = Model_name, color= 'salmon',linewidth=1) #Plot obs data on comparison figure   
    ax_compare[1].legend(loc='upper left')
    #ax_compare[1].set_xlabel ('Date'); 
    ax_compare[1].set_ylabel('Q ($m^3$/s)')

    Objective.index=Objective.index.astype('int')
    ax_compare[2].plot(Objective.index,Objective.iloc[:,0],marker='o', markersize=5, label = Model_name, linestyle='dashed', linewidth=0.5)
    ax_compare[2].plot ([Objective.index[0],Objective.index[-1]],[KGE_NWM,KGE_NWM], label = "NWM 2.1", color= 'forestgreen',linewidth=2.5)
    ax_compare[2].set_xlabel ('Iteration number')   
    ax_compare[2].set_ylabel('KGE')
    ax_compare[2].legend(loc='upper left')
    ax_compare[2].set_xlim([Objective.index[0]-1,Objective.index[-1]+1])
    # print(f'KGE_NWM: {KGE_NWM}')
    min_lim=min(KGE_NWM,0)
    ax_compare[2].set_ylim([min_lim,1])
    #ax_compare[2].set_ylim([Objective.iloc[:,0].min(),Objective.iloc[:,0].max()])            
    plt.figtext(0.05,0.64,"(a)",fontsize=14)
    plt.figtext(0.05,0.39,"(b)",fontsize=14)
    plt.figtext(0.05,0.08,"(c)",fontsize=14)
    fig.savefig(output_figure, bbox_inches='tight',dpi=300)     
    #plt.close(fig)

def Generate_Scatter_plot2(Model_name,Q_cms,Runoff_m3s,Next_Total_out,NWM_21,Objective,total_area,output_figure): 
    import sys

    from hydrotools.metrics import metrics 
    import seaborn as sns
    import numpy as np
    from math import sqrt
    sns.set_style("whitegrid")
    
   # Total_daily=Total.resample('D').apply(lambda x: np.sum(x.values))
    #if(len(NWM_21)>0):
   #     NWM_21_daily=NWM_21.resample('D').apply(lambda x: np.sum(x.values))
   # else:
   #     NWM_21_daily=pd.DataFrame()
    Q_cmh=Q_cms*3.600/(total_area*1000)
    Total_month=Q_cmh.resample('M').sum()
    
  
    Q_cmsNoNAN=Q_cms.dropna()
    #Nash=OB.nash_sutcliffe(Total_dailyNoNAN[Var_name_out],Total_dailyNoNAN['Obs_Q'])
    Nash=metrics.nash_sutcliffe_efficiency(Q_cmsNoNAN['q_cms_obs'], Q_cmsNoNAN['q_cms_sim'])
    NNash=1./(2.-Nash) 

    LogNash=metrics.nash_sutcliffe_efficiency(Q_cmsNoNAN['q_cms_obs'], Q_cmsNoNAN['q_cms_sim'],log=True)
    #LogNNash=1./(2.-LogNash) 
    
    KGE=metrics.kling_gupta_efficiency(Q_cmsNoNAN['q_cms_obs'], Q_cmsNoNAN['q_cms_sim'])
        
    select=Q_cmsNoNAN['q_cms_obs']>=Q_cmsNoNAN['q_cms_obs'].quantile(0.7)
    Temp=metrics.nash_sutcliffe_efficiency(Q_cmsNoNAN['q_cms_obs'][select], Q_cmsNoNAN['q_cms_sim'][select])
    NNash70=1./(2.-Temp) 

    select=Q_cmsNoNAN['q_cms_obs']<=Q_cmsNoNAN['q_cms_obs'].quantile(0.3)
    Temp=metrics.nash_sutcliffe_efficiency(Q_cmsNoNAN['q_cms_obs'][select], Q_cmsNoNAN['q_cms_sim'][select])
    NNash30=1./(2.-Temp) 
    
    TotQVolError=100.*(Q_cmsNoNAN['q_cms_sim'].sum()-Q_cmsNoNAN['q_cms_obs'].sum())/Q_cmsNoNAN['q_cms_obs'].sum()
    
    flood_criteria=Q_cmsNoNAN['q_cms_obs'].quantile(0.7)
    Q_cmsNoNAN['simulated_flood'] = (Q_cmsNoNAN['q_cms_sim'] >= flood_criteria)
    Q_cmsNoNAN['observed_flood'] = (Q_cmsNoNAN['q_cms_obs'] >= flood_criteria)
    
    contingency_table = metrics.compute_contingency_table(Q_cmsNoNAN['observed_flood'],Q_cmsNoNAN['simulated_flood'])                   
    TS = metrics.threat_score(contingency_table)
    coefOfpers = metrics.coefficient_of_persistence(Q_cmsNoNAN['q_cms_obs'],Q_cmsNoNAN['q_cms_sim']) 
    if(len(NWM_21)>0):  
        NWM_21NaN=NWM_21.loc[Q_cmsNoNAN.index]
        Nash_NWM=metrics.nash_sutcliffe_efficiency(Q_cmsNoNAN['q_cms_obs'], NWM_21NaN['flow_cms'])
        NNash_NWM=1./(2.-Nash_NWM) 
        KGE_NWM=metrics.kling_gupta_efficiency(Q_cmsNoNAN['q_cms_obs'], NWM_21NaN['flow_cms'])
    else: 
        Nash_NWM=""
        KGE_NWM=metrics.kling_gupta_efficiency(Q_cmsNoNAN['q_cms_obs'], NWM_21NaN['flow_cms'])
    
    mean_bias=metrics.mean_error(Q_cmsNoNAN['q_cms_obs'], Q_cmsNoNAN['q_cms_sim'])
    mean_absolute_error=OB.mean_absolute_error(Q_cmsNoNAN['q_cms_obs'], Q_cmsNoNAN['q_cms_sim'])
    
    RMSE=metrics.root_mean_squared_error(Q_cmsNoNAN['q_cms_obs'], Q_cmsNoNAN['q_cms_sim'])
    NRMSE = m.sqrt(RMSE)/Q_cmsNoNAN['q_cms_obs'].mean()
    NRMSE=OB.NRMSE(Q_cmsNoNAN['q_cms_sim'],Q_cmsNoNAN['q_cms_obs'])

    ########### Create plot
    # if output_figure = False, then do not create plot
    if output_figure != False:
        num_subplots = 2 #If plotting cumulative precipitation/soil moisture, use 3 subplots (flow, cumulative rain, soil moisture)
        plt.rcParams["figure.figsize"] = (4,8)
        fig, ax_compare = plt.subplots(num_subplots, 1) #Create the plot for this zone
    
        ax_compare[0].scatter(Q_cmsNoNAN['q_cms_obs'],Q_cmsNoNAN['q_cms_sim'], label = Model_name,s=13,edgecolors='black',color='silver',linewidths=0.3)
        if(len(NWM_21)>0): ax_compare[0].scatter(Q_cmsNoNAN['q_cms_obs'],NWM_21NaN['flow_cms'], label = 'NWM2.1',s=13,edgecolors='forestgreen',facecolors='none',linewidths=0.5,alpha=0.5)
        ax_compare[0].set_xlabel ('Obs Q ($m^3$/s)')
        ax_compare[0].set_ylabel('Sim Q ($m^3$/s)')
        ax_compare[0].set_aspect('equal',adjustable='box')
        ax_compare[0].legend(loc='upper left')
        
        max_y=1.05*max(Q_cmsNoNAN['q_cms_obs'].max(),Q_cmsNoNAN['q_cms_sim'].max())
        ax_compare[0].set_xlim([0,max_y])
        ax_compare[0].set_ylim([0,max_y])
        ax_compare[0].plot ([0,max_y],[0,max_y],linestyle='dashed',linewidth=0.5, color= 'r') #Plot obs data on comparison figure    
        if(len(NWM_21)>0): 
            ax_compare[0].set_title("NNSE = " + str(round(NNash,2)) +"(NWM 2.1: "+ str(round(NNash_NWM,2))+")")
        else:
            ax_compare[0].set_title("NNSE = " + str(round(NNash,2)) +" - No NWM 2.1")
        Total_month['Month']=Total_month.index.month
        #Total_month=Total_month[['Month',Var_name_out,'Obs_Q']]
        ax_compare[1]
        Obs_Month=Total_month[['Month','q_cms_obs']]
        Obs_Month['Variable']='Obs'
        Obs_Month=Obs_Month.rename(columns={'q_cms_obs':"Q"})
        Sim_Month=Total_month[['Month','q_cms_sim']]
        Sim_Month['Variable']='Sim'
        Sim_Month=Sim_Month.rename(columns={'q_cms_sim':"Q"})
        Month_Q=pd.concat([Obs_Month,Sim_Month])
        sns.boxplot(x='Month',y="Q",hue='Variable',data=Month_Q,ax=ax_compare[1])
        ax_compare[1].set_xlabel ('Month'); 
        ax_compare[1].set_ylabel('Q (Sum m/month)')
        ax_compare[1].yaxis.grid(True)
        ax_compare[1].xaxis.grid(True)
        plt.figtext(0.005,0.5,"(d)",fontsize=14)
        plt.figtext(0.005,0.08,"(e)",fontsize=14)
        #ax_compare[1].set_aspect('equal',adjustable='box')
        # ax_compare[1].scatter(Total_daily['Obs_Q'],Total_daily[Var_name_out], label = 'Daily Observed versus Simulated Q(m/d)', color= 'b') #Plot obs data on comparison figure        
        # max_y=1.05*max(Total_daily[Var_name_out].max(),Total_daily['Obs_Q'].max())
        # ax_compare[1].set_xlim([0,max_y])
        # ax_compare[1].set_ylim([0,max_y])
        # ax_compare[1].plot ([0,max_y],[0,max_y],linestyle='dashed',linewidth=0.5,alpha=0.8, color= 'r') #Plot obs data on comparison figure
        
        
        #ax_compare[1].set_title("Q Seasonality ")
        fig.savefig(output_figure, bbox_inches='tight',dpi=300)  
        #plt.close(fig)
    return NNash,LogNash,NNash70,NNash30,TotQVolError,KGE,Nash,mean_bias,mean_absolute_error,RMSE,NRMSE,TS,coefOfpers      

######
def plot_map_bymodel(Selected_calibration,All_Stats,stats,title,output_figure,MinModels,MinThreshold):
    import math as m
    Model_RR=['NWM2.1',"CFE","CFE_X","Topmodel"]

    if(stats=='KGE') | (stats=='NNash'):

        Selected_calibration['Best']=''
        Selected_calibration['Dif']=-9
        for i in range (0,len(Selected_calibration)):

            hru_id=Selected_calibration.index[i]
            Temp=All_Stats[All_Stats.index==hru_id]

            if(len(Temp)>=MinModels):
                Temp=Temp.sort_values(by=[stats],ascending=False)
                Selected_calibration.at[hru_id,'Best']=Temp.iloc[0]['Model_RR']
                if(Temp.iloc[0]['Model_RR']=='NWM2.1'):
                    Selected_calibration.at[hru_id,'Dif']=abs(Temp.iloc[0][stats]-Temp.iloc[1:4][stats].max())

                else:
                    Selected_calibration.at[hru_id,'Dif']=abs(Temp.iloc[0][stats]-Temp[Temp['Model_RR']=='NWM2.1'][stats].values[0])


    nmodels=len(Model_RR)
    nrow=m.ceil(nmodels/2)
    #figsize=(nrow*4, ncols*3)
    f, axes = plt.subplots(figsize=(16, 9.5), ncols=2, nrows=nrow)
    us_states_file="/home/west/us_states.json"
    us_states = gp.read_file(us_states_file)
    us_states=us_states[us_states.name != "Alaska"]
    us_states_bounds = us_states.geometry.total_bounds
    #contiguous_usa = gp.read_file(gplt.datasets.get_path('contiguous_usa'))    

    # countries[countries["name"] == "United States of America"].plot(color="lightgrey",
                                                 # ax=ax)

    #norm = mpl.colors.Normalize(vmin=0, vmax=2)
    #cmap = mpl.cm.ScalarMappable(norm=norm, cmap='bwr').cmap

    for i in range(0,len(Model_RR)):
        if(i==0): ax1=0;ax2=0;xtext=0.15;ytext=0.65
        if(i==1): ax1=0;ax2=1;xtext=0.55;ytext=0.65
        if(i==2): ax1=1;ax2=0;xtext=0.15;ytext=0.4
        if(i==3): ax1=1;ax2=1;xtext=0.55;ytext=0.4
        if(i==4): ax1=2;ax2=0;xtext=0.15;ytext=0.15
        if(i==5): ax1=2;ax2=1;xtext=0.55;ytext=0.15

        us_states.plot(ax=axes[ax1][ax2],color='silver', edgecolor='white')
        #Hyd_Sig.plot(ax=axes[0][0],column='Min_col',cmap='RdYlBu',categorical=True,legend="True" )

        All_Models_temp=All_Stats[All_Stats['Model_RR']==Model_RR[i]]
        CAMELS_516_temp=Selected_calibration[(Selected_calibration['Best']==Model_RR[i]) & (Selected_calibration['Dif']>=MinThreshold)]

        #if(ax2==1):
        # if(ax2==0): 
        # im=All_Models_temp.plot.scatter(x='Long',y='Lat',s=60,
        #            c=stats,colormap='bwr',vmin=0,vmax=1,
        #            title=Model_RR[i],ax=axes[ax1][ax2],colorbar=False); 

        im=axes[ax1][ax2].scatter(All_Models_temp['Long'],All_Models_temp['Lat'],s=60,
                   c=All_Models_temp[stats],vmin=0,vmax=1,cmap='Spectral');
        if(Model_RR[i]=="Topmodel"): Model_RR_str="TOPMODEL"
        else: Model_RR_str=Model_RR[i]
        axes[ax1][ax2].set_title(Model_RR_str, fontsize=13,weight='bold')
        # else:
        #     im=All_Models_temp.plot.scatter(x='Long',y='Lat',s=60,
        #         c=stats,colormap='bwr',vmin=0,vmax=1,
        #         title=Model_RR[i],ax=axes[ax1][ax2]);  
        axes[ax1][ax2].scatter(CAMELS_516_temp['Long'],CAMELS_516_temp['Lat'],s=110, facecolors='none', edgecolors='black',linewidth=2)

        #All_Models_temp=All_Models_temp.loc[['10244950','05591550','10336660','09306242']]
        #axes[ax1][ax2].scatter(All_Models_temp['Long'],All_Models_temp['Lat'],s=350, facecolors='none', edgecolors='green',linewidth=2)


        axes[ax1][ax2].set_yticklabels([])
        axes[ax1][ax2].set_xticklabels([])
        axes[ax1][ax2].set_xlim(([-125,-65]))
        axes[ax1][ax2].set_ylim(([25,50]))
        axes[ax1][ax2].xaxis.label.set_visible(False)
        axes[ax1][ax2].yaxis.label.set_visible(False)




        #plt.title(models[i])
        #plt.figtext(xtext,ytext,models[i],fontsize=10)

    plt.figtext(0.1,0.55,"(a)",fontsize=15)
    plt.figtext(0.52,0.55,"(b)",fontsize=15)
    plt.figtext(0.1,0.15,"(c)",fontsize=15)
    plt.figtext(0.52,0.15,"(d)",fontsize=15)
    cbar_ax = f.add_axes([0.93, 0.15, 0.02, 0.7])
    f.colorbar(im, cax=cbar_ax).set_label(label=stats,size=15,weight='bold')
    f.savefig(output_figure, bbox_inches='tight',dpi=300)


#####
 
def plot_box(Selected_calibration,All_Stats,stats,title,output_figure,MinModels):
    import math as m
    import seaborn as sns
    Model_RR=['NWM2.1',"CFE","CFE_X","Topmodel"]

    if(stats=='KGE') | (stats=='NNash'):
        All_Stats_Best=All_Stats.copy()
        Selected_calibration['Best']=''
        for i in range (0,len(Selected_calibration)):  
            
            hru_id=Selected_calibration.index[i]
            Temp=All_Stats[All_Stats.index==hru_id]
            
            if(len(Temp)>=MinModels):
                Temp=Temp.sort_values(by=[stats],ascending=False)
                Selected_calibration.at[hru_id,'Best']=Temp.iloc[0]['Model_RR']
      
                Temp_2=pd.DataFrame(data=Temp.iloc[0])
                Temp_2.at['Model_RR',hru_id]='Best'
                All_Stats_Best=pd.concat([All_Stats_Best,Temp_2.transpose()])
   
    #figsize=(nrow*4, ncols*3)
    Model_RR_Best=['NWM2.1',"CFE","CFE_X","TOPMODEL","Best"]
    All_Stats_Best['Model_RR'][All_Stats_Best['Model_RR']=="Topmodel"]="TOPMODEL"
    fig, ax = plt.subplots(figsize=(7.5, 4.5)) #Create the plot for this zone
    #fig.set_size_inches(0.8*len(Daily_Q_all_Model['Source'].unique()),4)
    
    All_Stats_Best[stats][All_Stats_Best[stats]<-1]=-1
   
    Line_75=All_Stats_Best[stats][All_Stats_Best['Model_RR']=="NWM2.1"].quantile(0.75)
    Line_50=All_Stats_Best[stats][All_Stats_Best['Model_RR']=="NWM2.1"].quantile(0.5)
    Line_25=All_Stats_Best[stats][All_Stats_Best['Model_RR']=="NWM2.1"].quantile(0.25)
    
    
    #plt.rcParams["figure.figsize"] = (5,10)
    sns.boxplot(x='Model_RR',y=stats,data=All_Stats_Best,ax=ax,color='silver',width=0.5,order=Model_RR_Best)
    #ax2 = ax.twinx()
    ax.hlines(Line_25, -0.5, 4.5,label='NWM 2.1 25th Percentile',linestyle="dashed", color ="orange")
    ax.hlines(Line_50, -0.5, 4.5,label='NWM 2.1 50th Percentile',linestyle="dashed", color ="green")
    ax.hlines(Line_75, -0.5, 4.5,label='NWM 2.1 75th Percentile',linestyle="dashed", color ="blue")
    fig.legend(loc = "lower center",ncol=3 )
    ax.set_xlabel ('',fontsize=14)
    ax.set_ylabel(stats)
    #ax.set_xticklabels(Daily_Q_all_Model['Source'].unique(),rotation=90)
    ax.yaxis.grid(True)
    #ax.set_yscale("log")
    #ax.set_title(Title_str)
    fig.savefig(output_figure, bbox_inches='tight',dpi=300)        
    #plt.close(fig)
    
    
def NumberOfModels_plot(Rainfall_Runoff_ModelAr,stats_name,stats_name_axis,stats_criteria,Selected_calibration,All_Stats,stats,title,output_figure,MinModels):

    CB_color_cycle = [ 
                  '#f781bf', '#a65628', '#984ea3',
                  '#999999', '#e41a1c', '#dede00','#377eb8','#ff7f00', '#4daf4a']
    
    Temp=Rainfall_Runoff_ModelAr.copy()
    Temp.append("NWM2.1")
    NumberOfModels=pd.DataFrame(index=Temp)
    for j in range(0,len(stats_name)):
        stats=stats_name[j]
        NumberOfModels[stats]=0
        
        for i in range (0,len(Selected_calibration)):  
            
            hru_id=Selected_calibration.index[i]
            Temp=All_Stats[All_Stats.index==hru_id]
            
            if(len(Temp)>=MinModels):
                if(stats_criteria[j]=="max"):
                    Temp=Temp.sort_values(by=[stats],ascending=False)
                    NumberOfModels.at[Temp.iloc[0]['Model_RR'],stats]=NumberOfModels.at[Temp.iloc[0]['Model_RR'],stats]+1
                elif(stats_criteria[j]=="min"):
                    Temp=Temp.sort_values(by=[stats],ascending=True)
                    NumberOfModels.at[Temp.iloc[0]['Model_RR'],stats]=NumberOfModels.at[Temp.iloc[0]['Model_RR'],stats]+1    
                else:  
                    Temp[stats]=np.abs(Temp[stats])
                    Temp=Temp.sort_values(by=[stats],ascending=True)
                    NumberOfModels.at[Temp.iloc[0]['Model_RR'],stats]=NumberOfModels.at[Temp.iloc[0]['Model_RR'],stats]+1
                    # print(stats)
                    # print(hru_id)
                    # print(Temp.iloc[0]['Model_RR'])
                    
    for i in range(0,len(stats_name)):
        NumberOfModels=NumberOfModels.rename(columns={stats_name[i]:stats_name_axis[i]})
    NumberOfModels.index=NumberOfModels.index.str.upper()
    
    if output_figure != False:
        fig, ax = plt.subplots() #Create the plot for this zone
        NumberOfModels.transpose().plot(kind="bar", stacked=True,ax=ax, colormap='BrBG')
        plt.title(title)
        ax.legend(loc='upper center',ncol=4)
        NMax=NumberOfModels.sum().max()
        ax.set_ylim([0,NMax+4])
        ax.set_xlabel ("Statistics")
        ax.set_ylabel ("Number of locations where model is best")
        fig.savefig(output_figure, bbox_inches='tight',dpi=300) 
    return NumberOfModels
    #plt.close(fig)


#########################################################################
# functions above here
# processes below
#######################################################################

# get date and use to label log file
current_date = datetime.now().strftime("%Y-%m-%d %H:%M")
cwd = os.getcwd()
log_file = f'{cwd}/plot_calibration_{current_date}.log'
# create log file
with open(log_file, 'w') as file:
    pass

CB_color_cycle = ['#377eb8', '#ff7f00', '#4daf4a',
          '#f781bf', '#a65628', '#984ea3',
          '#999999', '#e41a1c', '#dede00']

# BChoat; if output already produced for an HRU, should it be overwritten?
# if True, then checks if a sample file name exists and assume it indicates
# that all desired files exist
# overwrite_previous = False

# directory
Hydrofabrics_folder="/home/west/Projects/CAMELS/CAMELS_Files_Ngen/"
# topmodel outputs are placed in different folder since I moved CFE so I wouldn't mess it up
Hydrofabrics_folder_HD_tm="/media/Expansion/Projects/CAMELS/CAMELS_Files_Ngen/"
# Hydrofabrics_folder_HD_tm="/media/Expansion/Projects/CAMELS/CAMELS_Files_Ngen_Original/"

# directory holding usgs discharge data
dir_usgs = '/media/Expansion/Projects/CAMELS/USGS_Q_20071001_20131001'
# directory holding nwm 2.1 results
dir_nwm = '/media/Expansion/Projects/CAMELS/nwm_v2.1_chrt_camels513_csv/csv'

Hydrofabrics_folder_HD_cfe="/media/Expansion/Projects/CAMELS/CAMELS_Files_Ngen_Original/"
Select_calibration_file="/home/west/Projects/CAMELS/CAMELS_Files_Ngen/CAMELS_v3_calib_BestModelsAll.csv"
ALL=pd.read_csv(Select_calibration_file,dtype={'hru_id_CAMELS': str})
ALL=ALL.set_index(['hru_id_CAMELS'])
ALL=ALL.rename(columns={'N_Nexus':'N_nexus'})
ALL=ALL.drop(['Q_Best_Topmodel','Q_Best_CFE_X'],axis=1).dropna()

# BChoat 2024/01/26: Subset to only those catchments in the new topmodel
# output folder on expansion drive
# get hru_ids for hrus already reran with topmodel
files_tm = [f for f in os.listdir(f'{Hydrofabrics_folder_HD_tm}Results')]
files_tm = [f for f in files_tm if 'camels_' in f]
# print(files_tm)

hru_tm = [x.split('_')[1] for x in files_tm]

# print(f'hru_tm: {hru_tm}')
ALL = ALL[ALL.index.isin(hru_tm)]
# print(f'ALL: {ALL}')
Selected_calibration=ALL.copy()

# Selected_calibration['Lat']=-9.0
# Selected_calibration['Long']=-9.0
# for i in range (0,len(Selected_calibration)):  
#     hru_id=Selected_calibration.index[i]
#     Folder=Selected_calibration.iloc[i]['Folder_CAMELS']
#     Hydrofabrics=Hydrofabrics_folder+"/"+Folder+"/spatial/catchment_data.geojson"
#     basin = gp.read_file(Hydrofabrics) 
#     basin = basin.to_crs({'init':'epsg:4269'})
#     basin_id=basin.sort_values(by="area_sqkm",ascending=True).iloc[0].id
#     nexus_id=basin.sort_values(by="area_sqkm",ascending=True).iloc[0].toid
#     Selected_calibration.at[hru_id,'Lat']=basin.sort_values(by="area_sqkm",ascending=True).iloc[0].geometry.centroid.y
#     Selected_calibration.at[hru_id,'Long']=basin.sort_values(by="area_sqkm",ascending=True).iloc[0].geometry.centroid.x      
# directory
# NextGen_folder="/home/west/git_repositories/ngen_11112022/ngen"
# NextGen_folder = '/home/west/git_repositories/ngen_20240111_calib/ngen'
NextGen_folder = '/home/west/git_repositories/ngen_20240327_calib/ngen'


os.environ['PATH'] += os.pathsep+"/home/west/TauDEMDependencies/mpich/mpich-install"
os.environ['PATH']=NextGen_folder+"/venv/bin:/home/west/anaconda3/bin:/home/west/anaconda3/condabin:/sbin:/bin:/usr/bin:/usr/local/bin:/snap/bin"
str_sub="source "+NextGen_folder+"/venv/bin/activate"
out=subprocess.call(str_sub,shell=True)  

Results_calibration=Hydrofabrics_folder_HD_tm+"Calib_plots_WithNoah/"
# Results_calibration="/media/Expansion/Projects/CAMELS/TESTINGPLOTS/Calib_plots_WithNoah/"



# directory
if not os.path.exists(Results_calibration): 
    print('creating calibration output directory')
    os.mkdir(Results_calibration) 
Rainfall_Runoff_ModelAr=["CFE","CFE_X","Topmodel"]
# Rainfall_Runoff_ModelAr=["Topmodel"]

#Model="NOAH-CFE"
eval_type="Streamflow"   
flag_Plot_cat=0

b_doc_name=Results_calibration+"Appendix_B_11222022.docx"

# filename of output file(s) a .csv and .xlsx will be saved as file_out
# name with .csv extension
file_out = f'{Results_calibration}NWM_performance_first.csv'

# try to read in file_out in case it was created in previous execution (so we 
# don't reprocess results that have already been processed)
# otherwise, create empty dataframe to add results to
try:
    All_Stats = pd.read_csv(f"{Results_calibration}/NWM_performance_first.csv",
                        dtype={'hru_id_CAMELS': 'string'},                       
                         ).drop_duplicates(
                                subset=['hru_id_CAMELS', 'Model']
                         ).reset_index(
                                 drop = True
                                 ).set_index('hru_id_CAMELS')


    # All_Stats = pd.read_csv(file_out, 
    #         dtype = {'hru_id_CAMELS': 'string'}) #,
#             index_col = 0)

    # subset Selected_calibration to only HRUs that have not been previously processed
    Selected_calibration = Selected_calibration[
            ~Selected_calibration.index.isin(All_Stats.index)
            ]
                
except:
    All_Stats=pd.DataFrame()


document = Document()
change_orientation2()
document.add_heading('Appendix B: Plot for calibrated basins ', 0)
p = document.add_paragraph()



for i in range(0,len(Rainfall_Runoff_ModelAr)):    
    Selected_calibration[Rainfall_Runoff_ModelAr[i]]=0

#Selected_calibration=Selected_calibration.loc[['02430085']]
#Selected_calibration=Selected_calibration[(Selected_calibration['frac_snow']<0.1) & (Best_PET.str.contains('PET')) & (Selected_calibration['N_nexus']>=2)]
#Selected_calibration=Selected_calibration.sort_values(by="NCat")
count=0

#Selected_calibration=Selected_calibration.loc[['02430085','11284400','14362250']]
#Selected_calibration=Selected_calibration.loc[['11284400']]
# Selected_calibration=Selected_calibration.loc[['01052500', '01031500']]
# Selected_calibration=Selected_calibration.loc[['01052500', '01031500', '01055000']]
# Selected_calibration=Selected_calibration.loc[['01052500', '01055000']]






for i in range(0,len(Selected_calibration)):
    Folder_CAMELS=Selected_calibration.iloc[i]['Folder_CAMELS']
    print (f'Processing: {Folder_CAMELS}')
    hru_id_long = Folder_CAMELS.split('_')[1:3]
    hru_id_long = f'{hru_id_long[0]}_{hru_id_long[1]}'
    # print(f'hru_id_long: {hru_id_long}')
    hru_id=Selected_calibration.index[i]    
    # directory
    # work_dir=Hydrofabrics_folder_HD+Folder_CAMELS+"/"
    # Obs_Q_file=work_dir+"Validation/usgs_hourly_flow_2007-2019_"+hru_id+".csv" 
    # Obs_Q_file = f'{work_dir}Validation/usgs_hourly_flow_calibration.csv'
    Obs_Q_file = f'{dir_usgs}/usgs_hourly_flow_calibration_{hru_id}.csv'
    # NWM_21_file=work_dir+"Validation/nwm_v2.1_chrt."+hru_id+".csv"
    NWM_21_file= f'{dir_nwm}/nwm_v2.1_chrt.{hru_id}.csv'


    # BChoat 2024/01/26:
    # check if Obs_Q_file and NWM_21_file exists, 
    # if not just continue for now. Fix later. Looks like usgs data produced with 
    # /home/west/Projects/CAMELS/params_code/generate_validation_data.sh
    # 2024/05/05 Updated to saved observed data during calibration
    # if not os.path.exists(f'{work_dir}/Validation'):
    if not os.path.exists(Obs_Q_file):
        with open(log_file, 'a') as file:
            file.write(f'{Obs_Q_file} does not exist\n\n')

        continue
    #      os.makedirs(f'{work_dir}/Validation')
    # if not os.path.exists(Obs_Q_file):
    #     shutil.copy(f'{Hydrofabrics_folder_HD_cfe}/')

    
    
    # for j in range(0,3):
    for j in range(0, len(Rainfall_Runoff_ModelAr)):

        Rainfall_Runoff_Model=Rainfall_Runoff_ModelAr[j]
        # since I stored topmodel results in different location
        # define Hydrofabrics_folder_HD based on if CFE ot Topmodel
        if Rainfall_Runoff_Model == 'Topmodel':
            Hydrofabrics_folder_HD = Hydrofabrics_folder_HD_tm
            obj_log_name = "objective_log.txt"
            h5_name = "flowveldepth_Ngen.h5_last"
        else:
            Hydrofabrics_folder_HD = Hydrofabrics_folder_HD_cfe
            obj_log_name = "ngen-calibration_objective.txt"
            h5_name = "flowveldepth_Ngen.h5"
        
        # bchoat adding after editing to work of expansion
        work_dir = Hydrofabrics_folder_HD

        PET=Selected_calibration.iloc[i]['A_Best1']
        #PET="PET_4"
        if(hru_id=='02430085') & (hru_id=='02465493'):
            PET="PET_4"
        #PET="NOAH"
        os.chdir(Hydrofabrics_folder_HD)
        # os.chdir(work_dir)
        Model=PET+"_"+Rainfall_Runoff_Model
        Model_2="NOAH"+"_"+Rainfall_Runoff_Model
        calib_config_dir=work_dir+"/"+Model+"/"
        
        #Final_Folder=Hydrofabrics_folder_HD+"/Results/"+Folder_CAMELS+"/Calibration/"+Model  
        # directory 

        calib_config_dir=Hydrofabrics_folder_HD+"/Results/"+Folder_CAMELS+"/Calibration/"+Model+"/" +Model+"/"
        calib_config_dir_NOAH=Hydrofabrics_folder_HD+"/Results/"+Folder_CAMELS+"/Calibration/"+Model_2+"/" +Model_2+"/"
        # print(f'calib_config_dir: {calib_config_dir}\n..._NOAH: {calib_config_dir_NOAH}')


        # BChoat; edit so if defined plot already exists, then assume no need to produce further analysis,
        # so continue
        # if not overwrite_previous:
            # check get if plot already exists (assumes if the defined file exists, then no need for further analysis
#            test_file = Results_calibration+hru_id+"_"+Model+"_Scatter.png" 

            # print(f'line 598; test_file; {test_file}')        
#             if os.path.exists(test_file):
#                 continue
#            print("didn't pass due to test_file existing")
        
        Model_name=Model.replace("PET_","PET")
        Model_name=Model_name.replace("NOAH","NOM").upper()
        if(Selected_calibration['frac_snow'].iloc[i]>=0.1) & ("PET" in PET): Model_name="NOM_"+Model_name

############################################################################################################
# BChoat 2024/01/26
# I'm not sure any of the code between here and the next "#############" line far below is needed at this point
# between here and ~line 803
#############################################################################################################

        
        # if(not os.path.exists(Final_Folder)): os.makedirs(Final_Folder)
        # if(os.path.exists(calib_config_dir)):
        #     Run_mv="mv "+calib_config_dir+" " +Final_Folder
        #     out=subprocess.call(Run_mv,shell=True) 
        
        if(Rainfall_Runoff_Model=="CFE"): 
            calib_config_base=calib_config_dir+"/calib_config_CAMELS_CFE_Calib_Sep_2.yaml"
            iterations=300 
        elif(Rainfall_Runoff_Model=="CFE_X"): 
            calib_config_base=calib_config_dir+"/calib_config_CAMELS_CFE_X2_Calib_Sep_2.yaml"
            if(not os.path.exists(calib_config_base)):
                calib_config_base=calib_config_dir+"/calib_config_CAMELS_CFE_X_Calib_Sep_2.yaml"
            iterations=300 
        elif(Rainfall_Runoff_Model=="Topmodel"): 
            calib_config_base=calib_config_dir+"/calib_config_CAMELS_Topmodel_Calib_Sep_2.yaml"
            iterations=120 
        if('PET' in PET) & (Selected_calibration['frac_snow'].iloc[i]<0.1):  
            if(Rainfall_Runoff_Model=="CFE"): calib_realiz_base=calib_config_dir+"Realization_pet_cfe_calibNC.json"
            elif(Rainfall_Runoff_Model=="CFE_X"): calib_realiz_base=calib_config_dir+"Realization_pet_cfe_X_calibNC.json"
            elif(Rainfall_Runoff_Model=="Topmodel"): calib_realiz_base=calib_config_dir+"Realization_pet_Topmodel_calibNC.json"
            else: print("error: model does not exist")
            
        elif('PET' in PET) & (Selected_calibration['frac_snow'].iloc[i]>=0.1):
            if(Rainfall_Runoff_Model=="CFE"): calib_realiz_base=calib_config_dir+"Realization_noahowp_pet_cfe_calibNC.json"
            elif(Rainfall_Runoff_Model=="CFE_X"): calib_realiz_base=calib_config_dir+"Realization_noahowp_pet_cfe_X_calibNC.json"
            elif(Rainfall_Runoff_Model=="Topmodel"): calib_realiz_base=calib_config_dir+"Realization_noahowp_pet_Topmodel_calibNC.json"
            else: print("error: model does not exist")
          
        else:
           if(Rainfall_Runoff_Model=="CFE"): calib_realiz_base=calib_config_dir+"Realization_noahowp_cfe_calibNC.json"
           elif(Rainfall_Runoff_Model=="CFE_X"): calib_realiz_base=calib_config_dir+"Realization_noahowp_cfe_X_calibNC.json"
           elif(Rainfall_Runoff_Model=="Topmodel"): calib_realiz_base=calib_config_dir+"Realization_noahowp_Topmodel_calibNC.json"  
           else: print("error: model does not exist")
           PET="NOAH"        

        # since saved topmodel in new dir and diff name of objective file,
        # conditionally set file name here
        # objective_log = calib_config_dir+'/ngen-calibration_objective.txt' 
        # objective_log_2 = calib_config_dir_NOAH+'/ngen-calibration_objective.txt'  
        objective_log = f"{calib_config_dir}/{obj_log_name}"
        objective_log_2 = f"{calib_config_dir_NOAH}/{obj_log_name}"

        # if(os.path.exists(objective_log)):              
        #     os.makedirs(Hydrofabrics_folder_HD+"/Results/"+Folder_CAMELS+"/Calibration2/"+Model+"/")
        #     str_sub="mv "+calib_config_dir+" "+Hydrofabrics_folder_HD+"/Results/"+Folder_CAMELS+"/Calibration2/"+Model+"/"
        #     out=subprocess.call(str_sub,shell=True) 
        
        if( os.path.exists(objective_log)):
            if((os.path.exists(objective_log_2))):
                print ("BOTH AVALILABLE "+Folder_CAMELS)           

        if(not os.path.exists(objective_log)):
            ori_calib_config_dir=work_dir+"/"+Model+"/"+"/" + obj_log_name # 'ngen-calibration_objective.txt' 
            if((os.path.exists(objective_log))):
                print ("File Exists in original " + ori_calib_config_dir)                
                str_sub="mv "+work_dir+"/"+Model+" "+Hydrofabrics_folder_HD+"/Results/"+Folder_CAMELS+"/Calibration/"+Model+"/"
                out=subprocess.call(str_sub,shell=True) 
        size=0
        # print(f"line 668; objective_log:{objective_log}")
        if(os.path.exists(objective_log)):        
            size=os.path.getsize(objective_log)
        # print(f"line 670; size: {size}\nSelected_cali[N_nexus]: {Selected_calibration.iloc[i]['N_nexus']}")
        if(Selected_calibration.iloc[i]['N_nexus']>=2) & (size>0):
            # print(f'line 672; in if(....>=2) & size>0 ')
            flag_noFlow=0
            flag_NOAH_noFlow=0
            # print(Folder_CAMELS + " " + Model)
            #catchment_data_file = Base_directory/'spatial/catchment_data.geojson'
            #nexus_data_file = Base_directory/'spatial/nexus_data.geojson'
            # cross_walk_file  = work_dir+'parameters/cross-walk.json'
            cross_walk_file  = f'{Hydrofabrics_folder_HD_cfe}/camels_{hru_id_long}/parameters/cross-walk.json'

            # catchment_file=work_dir+'spatial/catchment_data.geojson'    
            catchment_file = f'{Hydrofabrics_folder_HD_cfe}/camels_{hru_id_long}/spatial/catchment_data.geojson'
            zones = gp.GeoDataFrame.from_file(catchment_file)    
            # nexus_file=work_dir+'spatial/nexus_data.geojson'    
            nexus_file=f'{Hydrofabrics_folder_HD_cfe}/camels_{hru_id_long}/spatial/nexus_data.geojson'    

            nexus = gp.GeoDataFrame.from_file(nexus_file)
            
            params_file  = calib_config_dir+'*.parquet'
            params_file = glob.glob(params_file)[0]

            # str_sub="mv ./Best_Results "+calib_config_dir
            # out=subprocess.call(str_sub,shell=True)
            

            
            output_file  = calib_config_dir+"/Best_Results/"+h5_name # 'flowveldepth_Ngen.h5' 
            output_file_NOAH  = calib_config_dir_NOAH+"/Best_Results/"+h5_name # 'flowveldepth_Ngen.h5' 
            
            ngen_objective  = calib_config_dir+obj_log_name # 'ngen-calibration_objective.txt'  
 
                 
            # Read objective_log
            Objective = pd.read_csv(ngen_objective,index_col=0,header=None)
            Indi_zero=Objective[Objective.iloc[:,0] == 0].index.values
            if(len(Indi_zero)==0):
                ind=0
            else:
                ind=Indi_zero[len(Indi_zero)-1]
            Objective=Objective.iloc[ind:,:]
            #Objective=Objective.reset_index()
            Objective=Objective[1].astype(float).to_frame()
            Diff_objective=Objective.diff(periods=1,axis=0)
            Diff_objective['Param']="-9"
            # Read calib_config_file file 
            try: # BChoat; adding try since topmodel saved in different location w/different 
                # ngen-cal version/different params
                with open(calib_config_base, 'r') as stream:
                    data_loaded = yaml.safe_load(stream)
                    # print(f"data_loaded: {data_loaded}")
                    start_t_calib=data_loaded['general']['evaluation_start']
                    end_t_calib = data_loaded['general']['evaluation_stop']        
                    n_iterations = data_loaded['general']['iterations']
            except:
                with open(calib_config_base, 'r') as stream:
                    data_loaded = yaml.safe_load(stream)
                    start_t_calib = data_loaded['model']['eval_params']['evaluation_start']
                    end_t_calib = data_loaded['model']['eval_params']['evaluation_stop']
                    n_iterations = data_loaded['general']['iterations']

 
            
            if(Objective.index[len(Objective)-1]>=n_iterations):
                
                Selected_calibration.at[hru_id,Rainfall_Runoff_ModelAr[j]]=1
                # Read realization file 
                # if(hru_id=='02430085') & (PET=="NOAH"):
                #    calib_realiz_base="/media/west/Expansion/Projects/CAMELS/CAMELS_Files_Ngen//Results/camels_02430085_18693151/Calibration/NOAH_CFE/NOAH_CFE/Realization_noahowp_cfe_calib_NCCsv.json"
                with open(calib_realiz_base) as fp:
                    data = json.load(fp)
                start_t_real = data['time']['start_time']
                end_t_real = data['time']['end_time']
                print(f'start_t_real: {start_t_real}')
                print(f'end_t_real: {end_t_real}')
        
                # Read calib_config_file file 
                # with open(calib_config_file, 'r') as stream:
                #     data_loaded = yaml.safe_load(stream)
                #     start_t_calib=data_loaded['general']['evaluation_start']
                #     end_t_calib = data['time']['end_time']
        
    
                # Read observed
            
                # Obs_Q_cms=pd.read_csv(Obs_Q_file,parse_dates=True,index_col=1)
                # Obs_Q_cms=Obs_Q_cms.loc[:,['q_cms']]
                # Obs_Q_cms.rename(columns = {'q_cms':'q_cms_obs'}, inplace = True)
                # BChoat; updating to reflect use of usgs Q data downloaded during calibration
                Obs_Q_cms=pd.read_csv(Obs_Q_file,parse_dates=True,index_col=0)
                # print(f'~line851; Obs_Q_cms.index: {Obs_Q_cms.index}')
                # print(Obs_Q_cms.columns)
                Obs_Q_cms=Obs_Q_cms.loc[:,['obs_flow']]
                Obs_Q_cms.rename(columns = {'obs_flow':'q_cms_obs'}, inplace = True)

                Obs_Q_cms = Obs_Q_cms[(Obs_Q_cms.index >= start_t_real) & (Obs_Q_cms.index <= end_t_real)]
                # print(f'~line857; Obs_Q_cms.index: {Obs_Q_cms.index}')

                
                # Read crosswalk
                data = json.load(open(cross_walk_file))
                Nexus=list(data.keys())
                for ii in range(0,len(Nexus)):
                    # print(f'cross-walk: {ii}')
                    data[Nexus[ii]].keys()
                    # print(f'{data[Nexus[ii]].keys()}')
                    if('Gage_no' in data[Nexus[ii]].keys()):
                        Nex=Nexus[ii]
                        # BChoat: Gauge_id is never actually used anywhere ... should it be?
                        Gauge_id=data[Nexus[ii]]['Gage_no'][0]
                        # print(f'Nex: {Nex}')
                        # print(f'Gauge_id: {Gauge_id}')

 
               # Read parameters            
                param=pd.read_parquet(params_file, engine='pyarrow')
                param=param.set_index(['param'])
                param=param.transpose()
                param=param.drop('model')
                param=param.drop('sigma')
                param=param.drop('0')
               
                param=param.astype(float)
                param_norm=param.copy()
                for j in range(0,len(param_norm.columns)):
                    param_norm[param_norm.columns[j]]=(param_norm[param_norm.columns[j]]-param_norm[param_norm.columns[j]].loc['min'])/(param_norm[param_norm.columns[j]].loc['max']-param_norm[param_norm.columns[j]].loc['min'])
                param_norm=param_norm.drop('max')
                param_norm=param_norm.drop('min')
                param=param.drop('max')
                param=param.drop('min')
               
                Objective=Objective.rename(columns={Objective.columns[0]:"Objective"})
                Objective.index=Objective.index.astype('str')
                param=pd.concat([param,Objective],axis=1)
                if(len(param.Objective.unique())<2):
                    # print('line 768; in if(lenParam.Objective.unique())<2)')
                    str_sub="rm -rf "+calib_config_dir
                    out=subprocess.call(str_sub,shell=True)
                else:
                    # print('line 772; in else()')
                    
                    Final_param=param.sort_values(by='Objective',ascending=False).iloc[0]
                    # print(f'output_file: {output_file}')        
                    if(not os.path.exists(output_file)):  
                       
                        # print("Did not find output, Need to re-run")
                        realiz_file=work_dir+"Realization_"+Model+"Best.json"
                        with open(calib_realiz_base) as fp:
                            data = json.load(fp) 
                        # print(f"data: {data.keys()}")
                        for i_p in range(0,len(Final_param)):
                            param_name=Final_param.index[i_p]
                            param_value=Final_param.iloc[i_p]
                            # BChoat; adding try since topmodel in diff location with different config file
                            # due to different version of ngen-cal being used
                            try:
                                if(param_name in data['global']['formulations'][0]['params']['modules'][1]['params']['model_params']):                           
                                    data['global']['formulations'][0]['params']['modules'][1]['params']['model_params'][param_name]=param_value
                                else:
                                    print("Key does not exist " + param_name)
                            except Exception as e:
                                print(e)
#                                 if(param_name in data['global']['formulations'][0]['params']['modules'][1]['params']['model_params']):                           
#                                     data['global']['formulations'][0]['params']['modules'][1]['params']['model_params'][param_name]=param_value
#                                 else:
#                                     print("Key does not exist " + param_name)


                        json_object=json.dumps(data,indent=4)
                        with open(realiz_file,"w") as outfile:
                            outfile.write(json_object)      
                        #define the ngen
                        os.chdir(work_dir)
                        str_sub="rm -rf ngen"
                        out=subprocess.call(str_sub,shell=True) 
                        str_sub="ln -s "+NextGen_folder
                        out=subprocess.call(str_sub,shell=True) 
                        print("Will run nextgen")
                        Run_nextgen="./ngen/cmake_build/ngen "+ "./spatial/catchment_data.geojson '' ./spatial/nexus_data.geojson '' "+realiz_file    
                        # if(len(zones)>8):
                        #     Run_nextgen="./ngen/cmake_build_paral/ngen "+ "./spatial/catchment_data.geojson '' ./spatial/nexus_data.geojson '' "+Realiz_out+ " " +hru_id+"_8.json"
                            
                        if("PET" in Model):
                            ID=Model.split("_")[1]
                            for idd in range(1,6):
                                Run_sed="sed -i 's/pet_method="+str(idd)+"/pet_method='"+str(ID)+"'/g' ./PET/*"
                                out=subprocess.call(Run_sed,shell=True)
                        out=subprocess.call(Run_nextgen,shell=True) 
                        
                        os.makedirs(calib_config_dir+"/Best_Results/", exist_ok=True)
                        output_file  = calib_config_dir+"/Best_Results/"+h5_name # 'flowveldepth_Ngen.h5' 
                        output_file_NOAH  = calib_config_dir_NOAH+"/Best_Results/"+h5_name # 'flowveldepth_Ngen.h5' 
                        #mv_files="mv flowveldepth_Ngen.h5 "+output_file
                        #out=subprocess.call(mv_files,shell=True) 
                        #mv_files="mv *.csv "+calib_config_dir+"/Best_Results/"
                        #out=subprocess.call(mv_files,shell=True) 
                      
# BChoat - Not sure if code above here and below prevous "#####" line is needed
####################################################################################                

                    # Read  simulated
                    if(os.path.exists(output_file)): 
                        Sim_Q_cms = pd.read_hdf(output_file)                       
                        Sim_Q_cms.index = Sim_Q_cms.index.map(lambda x: 'wb-'+str(x))
                        Sim_Q_cms = Sim_Q_cms[Sim_Q_cms.columns.drop(list(Sim_Q_cms.filter(regex='v')))]
                        Sim_Q_cms = Sim_Q_cms[Sim_Q_cms.columns.drop(list(Sim_Q_cms.filter(regex='d')))]
                        Sim_Q_cms=Sim_Q_cms.transpose().loc[:,[Nex]]
                        # print(f'Nex: {Nex}')
                        # print(f'Sim_Q_cms: \n{Sim_Q_cms.head(20)}')

                        Sim_Q_cms.rename(columns = {Nex:'q_cms_sim'}, inplace = True)
                        dt_range = pd.date_range(start_t_real,end_t_real, len(Sim_Q_cms)).round('min')
#                         print(f'dt_range: \n{dt_range}')

                        Sim_Q_cms.index=dt_range
                        # print(f'Sim_Q_cms: \n{Sim_Q_cms.head(20)}')
                        Sim_Q_cms_H = Sim_Q_cms.resample('1H').first()
                        # print(f'Sim_Q_cms_H:\n {Sim_Q_cms_H.head(20)}')
                        # BCHOAT ADDING FOR TROUBLESHOOTING
                        # print(os.getcwd())
                        # Sim_Q_cms.to_csv(f'{hru_id}_{Rainfall_Runoff_Model}.csv')
                        # Sim_Q_cms_H.to_csv(f'{hru_id}_{Rainfall_Runoff_Model}_H.csv')
                        tnx_file = glob.glob(calib_config_dir+"/Best_Results/"+"tnx-*.csv")[0]
                        tnx_df = pd.read_csv(tnx_file, index_col=0, parse_dates=[1], names=['ts', 'time', 'q_cms_sim']).set_index('time')
                        # print(f'~line985')
#                        print(f'Sim_Q_cms_H.index: {Sim_Q_cms_H.index}')
#                         print(f'tnx_df.index: {tnx_df.index}')
                        Sim_Q_cms_H=Sim_Q_cms_H['q_cms_sim']+tnx_df['q_cms_sim']
#                        Sim_Q_cms_H = Sim_Q_cms_H['q_cms_sim'] 
                        
                        if(os.path.exists(output_file_NOAH)): 
                            print(f'output_file NOAH exists {output_file_NOAH}')
                            Sim_Q_NOAH_cms = pd.read_hdf(output_file_NOAH)
                            Sim_Q_NOAH_cms.index = Sim_Q_NOAH_cms.index.map(lambda x: 'wb-'+str(x))
                            Sim_Q_NOAH_cms = Sim_Q_NOAH_cms[Sim_Q_NOAH_cms.columns.drop(list(Sim_Q_NOAH_cms.filter(regex='v')))]
                            Sim_Q_NOAH_cms = Sim_Q_NOAH_cms[Sim_Q_NOAH_cms.columns.drop(list(Sim_Q_NOAH_cms.filter(regex='d')))]
                            Sim_Q_NOAH_cms=Sim_Q_NOAH_cms.transpose().loc[:,[Nex]]
                            Sim_Q_NOAH_cms.rename(columns = {Nex:'q_cms_NOAH_sim'}, inplace = True)
                            dt_range = pd.date_range(start_t_real,end_t_real, len(Sim_Q_cms)).round('min')
                            Sim_Q_NOAH_cms.index=dt_range
                            Sim_Q_NOAH_cms_H = Sim_Q_NOAH_cms.resample('1H').first()
                            tnx_file = glob.glob(calib_config_dir+"/Best_Results/"+"tnx*.csv")[0]
                            tnx_df = pd.read_csv(tnx_file, index_col=0, parse_dates=[1], names=['ts', 'time', 'q_cms_NOAH_sim']).set_index('time')
                            Sim_Q_NOAH_cms_H=Sim_Q_NOAH_cms_H['q_cms_NOAH_sim']+tnx_df['q_cms_NOAH_sim']                    
#                             Sim_Q_NOAH_cms_H = Sim_Q_NOAH_cms_H['q_cms_NOAH_sim'] 

                        else:
                            Sim_Q_NOAH_cms_H=pd.DataFrame()
                            flag_NOAH_noFlow=1
                    else:
                        Sim_Q_cms_H=pd.DataFrame()
                        flag_noFlow=1
                    
                    
                    if(os.path.isfile(NWM_21_file)):
                        NWM_21=pd.read_csv(NWM_21_file,parse_dates=True,index_col=0)*100.
                    else: NWM_21=pd.DataFrame()
             
                    if(flag_noFlow==0): 
#                         print(' in flag_noFlow==0; ~line 1018')
                        # Q_cms = pd.concat([Sim_Q_cms_H, Obs_Q_cms], axis=1)
#                        print(f'~line1019; Sim_Q_cms_H.index: {Sim_Q_cms_H.index}')
#                        print(f'Obs_Q_cms.index: {Obs_Q_cms.index}')
                        Q_cms = pd.merge(Sim_Q_cms_H, Obs_Q_cms, 
                                            left_index = True, right_index = True,
                                            how = 'left')
                        if(flag_NOAH_noFlow==0): 
#                             print(' in flag_NOAH_noFlow==0; ~line 1024')

                            #Q_cms = pd.concat([Q_cms, Sim_Q_NOAH_cms_H], axis=1)
                            Q_cms = pd.merge(Sim_Q_cms_H, Obs_Q_cms, 
                                        left_index=True, right_index=True,
                                        how = 'left')
                    else:
#                         print('Defining Q_cms=Obs_Q_cms.copy(); ~line 1031')

                        Q_cms = Obs_Q_cms.copy()
                total_area=zones['area_sqkm'].sum()  
                if(flag_Plot_cat==1):
                    #Read  cat 
                      
                    Remove_cat=[]
                    icount=0
                    for index,row in zones.iterrows():
                        catstr=row[0]    
                        area=zones.area_sqkm.iloc[index]
                        
                        # READING OUTPUT CFE
                        
                        #if("tnx" not in zones.toid.iloc[index]):
                            
                        Cat_out_file=calib_config_dir+"/Best_Results/"+catstr+".csv"
                        Cat_out=pd.read_csv(Cat_out_file,parse_dates=True,index_col=1)
                        flag_to_outlet=0
                        #Cat_out=Cat_out.dropna()
                        if(isinstance(Cat_out.index.min(),str)):
                            Cat_out.index=pd.to_datetime(Cat_out.index)
                        
                        if(len(Cat_out.dropna())<len(Cat_out)):
                            #NAN_in_results.append([hru_id,Models[j],catstr,len(Cat_out.dropna()),len(Cat_out)])
                            print("NAN in " + hru_id + " " + catstr)
                        if(icount==0):
                            Total_out=Cat_out.copy()*area
                            #Rainfall_from_forcing_mmh_ave=Obs_RR_mmh_Temp.copy()*area
                        else:
                            Total_out=Total_out+Cat_out*area
                        icount=icount+1
                        #else:
                        #    Remove_cat.append(catstr)
                            #Rainfall_from_forcing_mmh_ave=Rainfall_from_forcing_mmh_ave+Obs_RR_mmh_Temp*area
                    # Convert m/h to m3/s - no need to multiply by area since I am alredy myltipying as I read it 
                    Runoff_m3s=Total_out['Q_OUT']*1000/3.6
                    
                    Total_out=Total_out/total_area
                    # Read  nex 
                                
                    for index,row in nexus.iterrows():
                        nexstr=row[0]  
                        # print(nexstr)
                        area=zones.area_sqkm.iloc[index]
                        
                        # READING OUTPUT CFE
                        Nex_out_file=calib_config_dir+"/Best_Results/"+nexstr+"_output.csv"
                        Nex_out=pd.read_csv(Nex_out_file,parse_dates=True,index_col=1,header=None)
                        
                        #Cat_out=Cat_out.dropna()
                        if(isinstance(Nex_out.index.min(),str)):
                            Nex_out.index=pd.to_datetime(Cat_out.index)
                        
                        if(len(Nex_out.dropna())<len(Nex_out)):
                            #NAN_in_results.append([hru_id,Models[j],catstr,len(Cat_out.dropna()),len(Cat_out)])
                            print("NAN in " + hru_id + " " + nexstr)
                        if(index==0):
                            Next_Total_out=Nex_out[2].copy()
                            #Rainfall_from_forcing_mmh_ave=Obs_RR_mmh_Temp.copy()*area
                        else:
                            Next_Total_out=Next_Total_out+Nex_out[2]
                            #Rainfall_from_forcing_mmh_ave=Rainfall_from_forcing_mmh_ave+Obs_RR_mmh_Temp*area                  
                else:
                    Next_Total_out=[]
                    Runoff_m3s=[]
                    
                output_figure=Results_calibration+hru_id+"_"+Model+"_TS.png"
                if(len(Runoff_m3s)>0): Runoff_m3s=Runoff_m3s[(Runoff_m3s.index>=start_t_calib) & (Runoff_m3s.index<=end_t_calib)]
                if(len(Next_Total_out)>0): Next_Total_out=Next_Total_out[(Next_Total_out.index>=start_t_calib) & (Next_Total_out.index<=end_t_calib)]
                # print(f'Q_cms.index: {Q_cms.index}')
                # print(f'start_t_calib: {start_t_calib}')
                # print(f'end_t_calib: {end_t_calib}')

                Q_cms=Q_cms[(Q_cms.index>=start_t_calib) & (Q_cms.index<end_t_calib)]
                # print(f'start_t_calib: {start_t_calib}')
                # print(f'end_t_calib: {end_t_calib}')
                NWM_21=NWM_21[(NWM_21.index>=start_t_calib) & (NWM_21.index<end_t_calib)]
                # print(Q_cms)
                if not os.path.exists(output_figure):
                    plot_calibration_1(Model_name,Q_cms,Runoff_m3s,Next_Total_out,NWM_21,Objective,total_area,output_figure)
                
                # Q_all = pd.merge(Q_cms,NWM_21, left_index=True, right_index=True)
                # Q_all = pd.concat([Q_cms, NWM_21], axis=1)
                # Q_all=Q_all.rename(columns={"flow_cms":"NWM21_cms"})
                # Q_all.to_csv(Results_calibration_Rachel+hru_id+"_"+Model+".csv")
                
                output_figure_2=Results_calibration+hru_id+"_"+Model+"_Scatter.png"
                # if os.path.exists(output_figure_2):
                #     output_figure_2 = False
                [NNash,LogNash,NNash70,NNash30,TotQVolError,KGE,Nash,mean_bias,mean_absolute_error,RMSE,NRMSE,TS,coefOfpers]=Generate_Scatter_plot2(Model_name,Q_cms,Runoff_m3s,Next_Total_out,NWM_21,Objective,total_area,output_figure_2)
                Temp=pd.DataFrame([[NNash,LogNash,NNash70,NNash30,TotQVolError,KGE,Nash,mean_bias,mean_absolute_error,RMSE,NRMSE,TS,coefOfpers,Selected_calibration.iloc[i]['Lat'],Selected_calibration.iloc[i]['Long']]],columns=['NNash','LogNash','NNash70','NNash30','TotQVolError','KGE','Nash','mean_bias','mean_absolute_error','RMSE','NRMSE','TS','coefOfpers','Lat','Long'])
                # print(f'Temp: {Temp}')
                Temp['hru_id_CAMELS']=hru_id
                Temp['Model']=Model
                Temp['Model_RR']=Rainfall_Runoff_Model
                Temp['Model_PET']=PET
                count=count+1
                text = "Formulation " + str(Model_name) + ", Basin: " + hru_id+" , NCat: "+str(Selected_calibration.iloc[i]['NCat'])+" \n"
                text = text+"aridity: "+str(round(Selected_calibration.iloc[i]['aridity'],2))+" - "
                text = text+"snow fraction: "+str(round(Selected_calibration.iloc[i]['frac_snow'],2))+" - "
                text = text+"seasonality: "+str(round(Selected_calibration.iloc[i]['p_seasonality'],2))
                
                p = document.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                r = p.add_run()
                # r.add_text('')

                r.add_picture(output_figure, width=Inches(5))
                r2 = p.add_run()
                r2.add_picture(output_figure_2, width=Inches(2.7))
                caption=": (a) Cumulative Q (m3), (b) Simulated and observed runoff (m/h), (c) KGE by DDS iteration, (d) Scatter plot of observed and simulated daily runoff versus simulated, and (e) seasonality of observed and simulated daily runoff"
                p = document.add_paragraph("Figure B."+str(count)+": "+text.replace("\n"," (").replace(" - ",", ")+")"+caption)
                
                # print('line 971!!!!!!!!!')
                All_Stats=pd.concat([All_Stats,Temp])


                NWM_df=All_Stats[(All_Stats.hru_id_CAMELS==hru_id) & (All_Stats.Model_RR=="NWM2.1")]
                
                if(len(NWM_df)==0):
 
                    Q_cmsNoNAN=Q_cms.dropna()
                    if(len(NWM_21)>0):  
                        NWM_21NaN=NWM_21.loc[Q_cmsNoNAN.index]
                        
                        Nash=metrics.nash_sutcliffe_efficiency(Q_cmsNoNAN['q_cms_obs'], NWM_21NaN['flow_cms'])
                        NNash=1./(2.-Nash) 
                        LogNash=metrics.nash_sutcliffe_efficiency(Q_cmsNoNAN['q_cms_obs'], NWM_21NaN['flow_cms'],log=True)
                        #LogNNash=1./(2.-LogNash) 

                        KGE=metrics.kling_gupta_efficiency(Q_cmsNoNAN['q_cms_obs'], NWM_21NaN['flow_cms'])
                        
                        select=Q_cmsNoNAN['q_cms_obs']>=Q_cmsNoNAN['q_cms_obs'].quantile(0.7)
                        Temp=metrics.nash_sutcliffe_efficiency(Q_cmsNoNAN['q_cms_obs'][select], NWM_21NaN['flow_cms'][select])
                        NNash70=1./(2.-Temp) 

                        select=Q_cmsNoNAN['q_cms_obs']<=Q_cmsNoNAN['q_cms_obs'].quantile(0.3)
                        Temp=metrics.nash_sutcliffe_efficiency(Q_cmsNoNAN['q_cms_obs'][select],NWM_21NaN['flow_cms'][select])
                        NNash30=1./(2.-Temp) 
                    
                        TotQVolError=100.*(NWM_21NaN['flow_cms'].sum()-Q_cmsNoNAN['q_cms_obs'].sum())/Q_cmsNoNAN['q_cms_obs'].sum()
                    

                        mean_bias=metrics.mean_error(Q_cmsNoNAN['q_cms_obs'], NWM_21NaN['flow_cms'])
                        mean_absolute_error=OB.mean_absolute_error(Q_cmsNoNAN['q_cms_obs'], NWM_21NaN['flow_cms'])
                    
                        RMSE=metrics.root_mean_squared_error(Q_cmsNoNAN['q_cms_obs'],NWM_21NaN['flow_cms'])
                        NRMSE = m.sqrt(RMSE)/Q_cmsNoNAN['q_cms_obs'].mean()
                        NRMSE=OB.NRMSE(NWM_21NaN['flow_cms'],Q_cmsNoNAN['q_cms_obs'])
                        flood_criteria=Q_cmsNoNAN['q_cms_obs'].quantile(0.7)
                        Q_cmsNoNAN['simulated_flood_NWM'] = (NWM_21NaN['flow_cms'] >= flood_criteria)
                        Q_cmsNoNAN['observed_flood_NWM'] = (Q_cmsNoNAN['q_cms_obs'] >= flood_criteria)
                        contingency_table = metrics.compute_contingency_table(Q_cmsNoNAN['observed_flood_NWM'],Q_cmsNoNAN['simulated_flood_NWM'])                   
                        TS = metrics.threat_score(contingency_table)
                        coefOfpers = metrics.coefficient_of_persistence(Q_cmsNoNAN['q_cms_obs'],NWM_21NaN['flow_cms']) 
                        
                    Temp=pd.DataFrame([[NNash,LogNash,NNash70,NNash30,TotQVolError,KGE,Nash,mean_bias,mean_absolute_error,RMSE,NRMSE,TS,coefOfpers,Selected_calibration.iloc[i]['Lat'],Selected_calibration.iloc[i]['Long']]],columns=['NNash','LogNash','NNash70','NNash30','TotQVolError','KGE','Nash','mean_bias','mean_absolute_error','RMSE','NRMSE','TS','coefOfpers','Lat','Long'])
                    Temp['hru_id_CAMELS']=hru_id
                    Temp['Model']="NWM2.1"
                    Temp['Model_RR']="NWM2.1"
                    Temp['Model_PET']="NWM2.1"  
                    All_Stats=pd.concat([All_Stats,Temp])
                     
                # output_figure=Results_calibration+hru_id+"_"+Model+"Params.png"
                # num_subplots = 1 #If plotting cumulative precipitation/soil moisture, use 3 subplots (flow, cumulative rain, soil moisture)
                # plt.rcParams["figure.figsize"] = (8,4)
                # fig, ax_compare = plt.subplots(num_subplots, 1) #Create the plot for this zone
                # #ax_compare[0].set_title('Cumulative Precipitation')
                # sns.stripplot(param_norm,ax=ax_compare)
                
                # ax_compare.set_ylabel('Normalized param')
                # ax_compare.set_xticklabels(param_norm.columns,rotation=45, ha='right')
                # #ax_compare[0].set_xlabel ('Date'); 
    
                # fig.savefig(output_figure, bbox_inches='tight',dpi=300)     
                # plt.close(fig)           

# print(f"b_doc_name: {b_doc_name}")
document.save(b_doc_name)
# print(f'line 1035: All_stats: {All_Stats.head()}')
                
All_Stats=All_Stats.set_index(['hru_id_CAMELS'])



stats_name=['KGE','TotQVolError','NNash','LogNash','mean_bias','NRMSE','TS','coefOfpers']
stats_criteria=['max','abs_min','max','max','min','min','max','max']
stats_name_axis=['KGE','TotQVolError','NNash','LogNash','mean_bias','NRMSE','TS-70th','coefOfpers']
#----------------4-----------------------------
stats="KGE"
output_figure=Results_calibration+"RR_"+stats+"map_4.png"
title="KGE - multiple models"

MinModels=4


All_Stats_4=All_Stats.copy()
drop=[]
Unique_index=All_Stats.index.unique()
for i in range (0,len(All_Stats.index.unique())):      
    hru_id=Unique_index[i]
    Temp=All_Stats[All_Stats.index==hru_id]
    if(len(Temp)<MinModels):
        drop.append(hru_id)
stats="KGE" 
All_Stats_4=All_Stats_4.drop(drop)
MinThreshold=0.0
output_figure=Results_calibration+"RR_"+stats+"map_"+str(MinModels)+"NoMarkup.png"
plot_map_bymodel(Selected_calibration,All_Stats_4,stats,title,output_figure,MinModels,MinThreshold)
MinThreshold=0.05
output_figure=Results_calibration+"RR_"+stats+"map_"+str(MinModels)+"NoMarkupTh_005.png"
plot_map_bymodel(Selected_calibration,All_Stats_4,stats,title,output_figure,MinModels,MinThreshold)

output_figure=Results_calibration+"BoxPlot"+stats+"_"+str(MinModels)+".png"
plot_box(Selected_calibration,All_Stats_4,stats,title,output_figure,MinModels)

stats="NNash"
output_figure=Results_calibration+"RR_"+stats+"map_"+str(MinModels)+"NoMarkup.png"
title="NNSE - multiple models"
MinThreshold=0
plot_map_bymodel(Selected_calibration,All_Stats_4,stats,title,output_figure,MinModels,MinThreshold)

output_figure=Results_calibration+"BoxPlot"+stats+"_"+str(MinModels)+".png"
plot_box(Selected_calibration,All_Stats_4,stats,title,output_figure,MinModels)

output_figure=Results_calibration+"BoxPlot"+stats+"_"+str(MinModels)+".png"
Selected_calibration=Selected_calibration.drop(['Best'], axis=1)
title="Best model by stats - "+str(MinModels)
output_figure=Results_calibration+"NumberOfModels"+"_"+str(MinModels)+".png"
NumberOfModels=NumberOfModels_plot(Rainfall_Runoff_ModelAr,stats_name,stats_name_axis,stats_criteria,Selected_calibration,All_Stats_4,stats,title,output_figure,MinModels)
from datetime import date
str(date.today())
import csv


Selected_calibration.to_csv(Results_calibration+"AllBasins"+str(date.today())+"1248.csv", quotechar='"',quoting=csv.QUOTE_NONNUMERIC)
Selected_calibration.to_excel(Results_calibration+"AllBasins"+str(date.today())+"1248.xlsx")


# save all output
All_Stats.to_csv(file_out)

xcl_out = file_out.replace('.csv', '.xlsx')

All_Stats.to_excel(xcl_out)
# BChoat editing to write to csv as well, so can append

